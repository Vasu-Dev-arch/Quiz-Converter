import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.oxml.ns import qn
import re

ns = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
}

def omml_to_text(elem):
    tag_name = elem.tag.split('}')[-1]
    if tag_name == 't':
        return elem.text or ''
    elif tag_name == 'sSup':
        base_elem = elem.find(qn('m:e'), namespaces=ns)
        sup_elem = elem.find(qn('m:sup'), namespaces=ns)
        base = omml_to_text(base_elem) if base_elem is not None else ''
        sup = omml_to_text(sup_elem) if sup_elem is not None else ''
        return base + '^' + sup
    elif tag_name == 'sSub':
        base_elem = elem.find(qn('m:e'), namespaces=ns)
        sub_elem = elem.find(qn('m:sub'), namespaces=ns)
        base = omml_to_text(base_elem) if base_elem is not None else ''
        sub = omml_to_text(sub_elem) if sub_elem is not None else ''
        return base + '_' + sub
    elif tag_name == 'sSubSup':
        base_elem = elem.find(qn('m:e'), namespaces=ns)
        sub_elem = elem.find(qn('m:sub'), namespaces=ns)
        sup_elem = elem.find(qn('m:sup'), namespaces=ns)
        base = omml_to_text(base_elem) if base_elem is not None else ''
        sub = omml_to_text(sub_elem) if sub_elem is not None else ''
        sup = omml_to_text(sup_elem) if sup_elem is not None else ''
        return base + '_' + sub + '^' + sup
    elif tag_name == 'frac':
        num_elem = elem.find(qn('m:num'), namespaces=ns)
        den_elem = elem.find(qn('m:den'), namespaces=ns)
        num = omml_to_text(num_elem) if num_elem is not None else ''
        den = omml_to_text(den_elem) if den_elem is not None else ''
        return '(' + num + '/' + den + ')'
    elif tag_name == 'rad':
        deg_elem = elem.find(qn('m:deg'), namespaces=ns)
        base_elem = elem.find(qn('m:e'), namespaces=ns)
        deg = omml_to_text(deg_elem) if deg_elem is not None else ''
        base = omml_to_text(base_elem) if base_elem is not None else ''
        if deg:
            return 'root^' + deg + '(' + base + ')'
        else:
            return '√(' + base + ')'
    elif tag_name == 'd':
        base_elem = elem.find(qn('m:e'), namespaces=ns)
        base = omml_to_text(base_elem) if base_elem is not None else ''
        return '(' + base + ')'
    elif tag_name == 'r':
        return ''.join(omml_to_text(child) for child in elem)
    else:
        return ''.join(omml_to_text(child) for child in elem if child is not None)

def get_para_text(para):
    parts = []
    for child in para._element.iterchildren():
        if child.tag == qn('w:r'):
            t = child.find(qn('w:t'))
            if t is not None:
                parts.append(t.text or '')
        elif child.tag == qn('m:oMath'):
            math_text = omml_to_text(child)
            parts.append(math_text)
    return ''.join(parts)

def process_block(block_lines):
    block_text = '\n'.join(block_lines)
    if 'Answer:' not in block_text or 'Options:' not in block_text:
        return None

    # Find question part
    question_line = block_lines[0]
    opts_match = re.search(r'Options:\s*(.*)', question_line, re.DOTALL | re.IGNORECASE)
    if opts_match:
        question_part = question_line[:opts_match.start()].strip()
        opts_part = opts_match.group(1).strip()
    else:
        question_part = question_line.strip()
        opts_part = None
        for j in range(1, len(block_lines)):
            if 'Options:' in block_lines[j]:
                opts_part = block_lines[j].split('Options:')[1].strip()
                break
        if opts_part is None:
            return None

    # Extract question, handling ** or not
    q_match = re.match(r'^\d+\.\s*\*\*(.*)\*\*', question_part)
    if q_match:
        question = q_match.group(1).strip()
    else:
        q_match = re.match(r'^\d+\.\s*(.*)', question_part)
        if q_match:
            question = q_match.group(1).strip()
        else:
            return None

    # Options
    option_texts = re.findall(r'\(([a-d])\)\s*(.+?)(?=\s*\([a-d]\)|$)', opts_part, re.DOTALL | re.IGNORECASE)
    options = ['', '', '', '']
    for letter, text in option_texts:
        idx = ord(letter.lower()) - ord('a')
        if 0 <= idx < 4:
            options[idx] = text.strip().replace('\n', ' ').strip()

    if not any(options):
        return None

    # Answer
    try:
        ans_line = next(l for l in block_lines if '**Answer:**' in l or 'Answer:' in l)
    except StopIteration:
        return None
    ans_match = re.search(r'\*\*Answer:\*\*\s*\(([a-d])\)\s*(.*)', ans_line, re.IGNORECASE)
    if not ans_match:
        ans_match = re.search(r'Answer:\s*\(([a-d])\)\s*(.*)', ans_line, re.IGNORECASE)
    if not ans_match:
        return None
    ans_letter = ans_match.group(1).lower()

    # Explanation
    exp_index = block_lines.index(ans_line) + 1
    exp_lines = [line.strip() for line in block_lines[exp_index:] if line.strip()]
    exp = ' '.join(exp_lines)
    exp = re.sub(r'\*\*Explanation:\*\*\s*', '', exp, flags=re.IGNORECASE)
    exp = re.sub(r'Explanation:\s*', '', exp, flags=re.IGNORECASE).strip()

    return {
        'question': question,
        'options': options,
        'ans_letter': ans_letter,
        'explanation': exp
    }

def browse_input():
    path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if path:
        input_var.set(path)

def browse_output():
    path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if path:
        output_var.set(path)

def convert():
    input_path = input_var.get()
    output_path = output_var.get()
    if not input_path or not output_path:
        messagebox.showerror("Error", "Please select input and output files.")
        return
    try:
        doc = Document(input_path)
        lines = []
        for p in doc.paragraphs:
            text = get_para_text(p)
            if text.strip():
                lines.append(text)
        questions = []
        current_block = []
        for line in lines:
            stripped = line.strip()
            if re.match(r'^[\u002d\u2013\u2014\u2015]+$', stripped) or stripped == '':
                if current_block and re.match(r'^\d+\.', current_block[0]):
                    q = process_block(current_block)
                    if q:
                        questions.append(q)
                current_block = []
            else:
                if re.match(r'^\d+\.', stripped):
                    if current_block:
                        q = process_block(current_block)
                        if q:
                            questions.append(q)
                    current_block = [line]
                elif current_block:
                    current_block[-1] += ' ' + line  # Append to last if continuation
                else:
                    # Skip headers
                    pass
        if current_block:
            q = process_block(current_block)
            if q:
                questions.append(q)
        if not questions:
            messagebox.showwarning("Warning", "No questions found in the document.")
            return
        out_doc = Document()
        for q in questions:
            table = out_doc.add_table(rows=8, cols=3)
            table.style = 'Table Grid'
            rows = table.rows
            # Row 0: Question
            rows[0].cells[0].text = 'Question'
            merged = rows[0].cells[1].merge(rows[0].cells[2])
            merged.text = q['question']
            # Row 1: Type
            rows[1].cells[0].text = 'Type'
            merged = rows[1].cells[1].merge(rows[1].cells[2])
            merged.text = 'multiple_choice'
            # Rows 2-5: Options
            for idx, opt in enumerate(q['options']):
                r = idx + 2
                rows[r].cells[0].text = 'Option'
                rows[r].cells[1].text = opt
                correct = 'correct' if chr(ord('a') + idx) == q['ans_letter'] else 'incorrect'
                rows[r].cells[2].text = correct
            # Row 6: Solution
            rows[6].cells[0].text = 'Solution'
            merged = rows[6].cells[1].merge(rows[6].cells[2])
            exp_parts = re.split(r'\s*\n\s*', q['explanation'])
            if exp_parts:
                p = merged.paragraphs[0]
                p.text = exp_parts[0]
                for part in exp_parts[1:]:
                    if part.strip():
                        merged.add_paragraph(part)
            # Row 7: Marks
            rows[7].cells[0].text = 'Marks'
            rows[7].cells[1].text = '1'
            rows[7].cells[2].text = '0'
            out_doc.add_paragraph('')
        out_doc.save(output_path)
        messagebox.showinfo("Success", f"Conversion completed. Processed {len(questions)} questions.")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

root = tk.Tk()
root.title("DOCX Question Converter")

input_var = tk.StringVar()
output_var = tk.StringVar()

tk.Label(root, text="Input File:").grid(row=0, column=0, padx=10, pady=5)
tk.Entry(root, textvariable=input_var, width=50).grid(row=0, column=1, padx=10, pady=5)
tk.Button(root, text="Browse", command=browse_input).grid(row=0, column=2, padx=10, pady=5)

tk.Label(root, text="Output File:").grid(row=1, column=0, padx=10, pady=5)
tk.Entry(root, textvariable=output_var, width=50).grid(row=1, column=1, padx=10, pady=5)
tk.Button(root, text="Browse", command=browse_output).grid(row=1, column=2, padx=10, pady=5)

tk.Button(root, text="Convert", command=convert).grid(row=2, column=1, pady=20)

root.mainloop()