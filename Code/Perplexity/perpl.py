import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
from docx import Document
from docx.shared import Pt
import re
import os

def merge_cells(row, start, end):
    row.cells[start].merge(row.cells[end])

def create_question_table(doc, qdata):
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'

    # Row 0: Question (merge col 1 & 2)
    table.cell(0, 0).text = "Question"
    cell_q = table.cell(0, 1)
    cell_q.text = qdata['question']
    merge_cells(table.rows[0], 1, 2)

    # Row 1: Type (merge col 1 & 2)
    table.cell(1, 0).text = "Type"
    cell_type = table.cell(1, 1)
    cell_type.text = "multiple_choice"
    merge_cells(table.rows[1], 1, 2)

    # Rows 2–5: Options (label, content, correct/incorrect)
    for i in range(4):
        table.cell(2 + i, 0).text = "Option"
        table.cell(2 + i, 1).text = qdata['options'][i] if i < len(qdata['options']) else ""
        table.cell(2 + i, 2).text = "correct" if i == qdata['correct'] else "incorrect"

    # Row 6: Solution (merge col 1 & 2)
    table.cell(6, 0).text = "Solution"
    cell_sol = table.cell(6, 1)
    cell_sol.text = qdata.get('solution', "")
    merge_cells(table.rows[6], 1, 2)

    # Row 7: Marks
    table.cell(7, 0).text = "Marks"
    table.cell(7, 1).text = "1"
    table.cell(7, 2).text = "0"

def find_answer_index(options, ans_text):
    ans_text = ans_text.strip().lower()
    # a, (a), b etc
    m = re.match(r'^[\(\[]?([a-d])[.\)\]]?', ans_text)
    if m:
        idx = ord(m.group(1)) - ord('a')
        if 0 <= idx < 4:
            return idx
    # Try to match answer text with option content (case-insensitive, strip spaces)
    for i, op in enumerate(options):
        if ans_text and ans_text in op.lower():
            return i
    return 0

def parse_docx_questions(filepath):
    doc = Document(filepath)
    text_blocks = []
    buf = []

    # Gather all non-empty paragraphs, ignore decorations
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        # Treat common separators as block splitters
        if re.fullmatch(r'-{2,}|—{1,}', t) or re.fullmatch(r'\\d+\\.', t):
            if buf:
                text_blocks.append(buf)
                buf=[]
        else:
            buf.append(t)
    if buf: text_blocks.append(buf)

    questions = []
    for block in text_blocks:
        question_lines = []
        options = []
        answer = ""
        solution = ""
        opt_stage = False
        solution_stage = False
        answer_stage = False

        for line in block:
            # Single line "Options:" with inline options e.g. "Options: (a)...(b)..."
            if not opt_stage and re.match(r'^Options?:', line, re.I):
                # Extract inline options, e.g. "Options: (a) foo (b) bar (c) baz (d) qux"
                opts_inline = re.findall(r'\\([a-d]\\)|([a-d])\\.\\s*([^\\(\\[]+?)(?=\\s*\\([a-d]\\)|\\b[a-d]\\.|\s*$)', line, re.I)
                if opts_inline:
                    # Each tuple may have either group 2 or 3, flatten
                    col_opts = []
                    for o in opts_inline:
                        for group in o[1:]:
                            if group: col_opts.append(group.strip())
                    options = col_opts if col_opts else options
                else:
                    # Or just start options stage, if not inline
                    opt_stage=True
            elif re.match(r'^[\\(\\[]?[a-d][.\\)\\]]? ', line):
                opt_stage = True

            if not opt_stage and not solution_stage and not answer_stage:
                # Not in option/solution/answer section: still accumulating question
                if 'Options:' in line:
                    before = line[:line.index('Options:')]
                    if before.strip(): question_lines.append(before.strip())
                else:
                    question_lines.append(line)
            # Option lines
            elif opt_stage and re.match(r'^[\\(\\[]?[a-d][.\\)\\]]? ', line):
                # Option lines: "a. XXX", "(b) YYY" etc.
                option_str = re.sub(r'^[\\(\\[]?[a-d][.\\)\\]]? ', '', line)
                options.append(option_str.strip())
            # Special handling if options finished and answer/solution start
            elif re.match(r'^(Answer|Ans|Correct):', line, re.I):
                answer_stage = True
                opt_stage=False
                answer = re.sub(r'^(Answer|Ans|Correct):\\s*', '', line, flags=re.I).strip()
            elif re.match(r'^(Explanation|Solution):', line, re.I):
                solution_stage=True
                sol_tmp = re.sub(r'^(Explanation|Solution):\\s*', '', line, flags=re.I).strip()
                if sol_tmp: solution = sol_tmp
            elif solution_stage:
                solution += (" " + line).strip()

        qtext = " ".join(question_lines).strip()
        # Remove trailing "Options:" if any
        qtext = re.sub(r'Options:.*', '', qtext, flags=re.I).strip()

        # Ensure four options, pad
        while len(options) < 4:
            options.append("")
        options = options[:4]
        correct_index = find_answer_index(options, answer) if answer else 0

        questions.append({
            'question': qtext,
            'options': options,
            'correct': correct_index,
            'solution': solution,
        })

    return questions

############# Modern minimal UI #############
class QuizFormatterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("QuizFormatter - Professional")
        self.root['bg'] = "#EFF2F7"
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("TButton", font=("Arial", 11))
        style.configure("TLabel", background="#EFF2F7", font=("Arial", 11))
        # App Header
        ttk.Label(root, text="QuizFormatter", font=('Arial', 22,'bold'), foreground="#2166ac").pack()
        ttk.Label(root, text="Upload a .docx and get structured quiz tables.", font=("Arial", 11)).pack(pady=(0,12))
        main = ttk.Frame(root, padding=18)
        main.pack(fill="both", expand=True)
        filefr = ttk.Frame(main)
        filefr.pack(fill='x')
        self.file_label = ttk.Label(filefr, text="No file selected", font=("Arial", 10))
        self.file_label.pack(side='left')
        self.upload_btn = ttk.Button(filefr, text="Upload .docx", command=self.load_file)
        self.upload_btn.pack(side='right')
        btn_frame = ttk.Frame(main)
        btn_frame.pack(pady=8)
        self.preview_btn = ttk.Button(btn_frame, text="Preview", state='disabled', command=self.preview)
        self.preview_btn.pack(side='left', padx=5)
        self.save_btn = ttk.Button(btn_frame, text="Convert & Save", state='disabled', command=self.convert_save)
        self.save_btn.pack(side='left', padx=5)
        self.console = ScrolledText(main, height=12, font=("Consolas", 10))
        self.console.pack(fill='both', expand=True, pady=(8,0))
        self.filepath = None
        self.questions = []

    def log(self, msg, tag=None):
        self.console.insert('end', msg + '\n')
        self.console.see('end')

    def load_file(self):
        fn = filedialog.askopenfilename(filetypes=[("Word .docx", "*.docx")])
        if not fn: return
        self.filepath = fn
        self.file_label.config(text=os.path.basename(fn))
        try:
            self.questions = parse_docx_questions(fn)
            if not self.questions:
                self.log("No questions found. Please check the file.")
                self.preview_btn.state(['disabled'])
                self.save_btn.state(['disabled'])
            else:
                self.log(f"{len(self.questions)} questions loaded.")
                self.preview_btn.state(['!disabled'])
                self.save_btn.state(['!disabled'])
        except Exception as e:
            self.log(f"Error: {e}")
            self.preview_btn.state(['disabled'])
            self.save_btn.state(['disabled'])

    def preview(self):
        if not self.questions:
            messagebox.showinfo("No Data", "No questions found.")
            return
        w = tk.Toplevel(self.root)
        w.title("Preview")
        w.geometry("850x600")
        st = ScrolledText(w, font=("Arial", 11))
        st.pack(fill='both', expand=True)
        for i, q in enumerate(self.questions, 1):
            st.insert('end', f"Q{i}: {q['question']}\n")
            for oi,opt in enumerate(q['options']):
                mark = " (correct)" if oi==q['correct'] else ""
                st.insert('end', f"   {chr(97+oi)}. {opt}{mark}\n")
            st.insert('end', f"Solution: {q['solution']}\n{'-'*60}\n\n")
        st['state'] = 'disabled'

    def convert_save(self):
        if not self.questions:
            messagebox.showerror("No Data", "No questions found to convert!")
            return
        fn = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word .docx", "*.docx")])
        if not fn: return
        doc = Document()
        for q in self.questions:
            create_question_table(doc, q)
            doc.add_paragraph()
        doc.save(fn)
        self.log(f"Document saved: {fn}")
        messagebox.showinfo("Done!", f"Quiz document saved:\n{fn}")

def main():
    root = tk.Tk()
    app = QuizFormatterApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
