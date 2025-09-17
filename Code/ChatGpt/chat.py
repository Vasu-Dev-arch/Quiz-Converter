# quiz_formatter_app_windows_ready.py
# PySide6 desktop UI — Windows-ready, dark-theme default, true pill buttons, aligned inputs/buttons.
# pip install PySide6 python-docx lxml

import sys, re, os, unicodedata
from pathlib import Path
from lxml import etree
from docx import Document
from PySide6.QtCore import Qt, QTimer
from PySide6.QtGui import QFont, QKeySequence, QAction
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QLineEdit, QPushButton,
    QHBoxLayout, QVBoxLayout, QFileDialog, QMessageBox, QFrame, QSizePolicy
)

# --- Conversion Logic Start --- (starts at line 28, ends at line 289 in this file)

# XML namespaces for docx parsing
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
}

# ----------------------
# Low-level helpers
# ----------------------

def paragraph_full_text(paragraph):
    """
    Reconstruct paragraph text by iterating runs and math nodes in order.
    This preserves OMath content (extracts m:t text) and normal text.
    """
    xml = paragraph._p.xml.encode("utf-8")
    root = etree.fromstring(xml)
    pieces = []
    for child in root:
        tag = etree.QName(child.tag).localname
        if tag == "r":  # run
            for t in child.findall('.//w:t', namespaces=NSMAP):
                if t.text:
                    pieces.append(t.text)
            # preserve explicit breaks
            if child.findall('.//w:br', namespaces=NSMAP):
                pieces.append('\n')
        elif tag in ("oMath", "oMathPara"):
            # gather math text tokens
            for mt in child.findall('.//m:t', namespaces=NSMAP):
                if mt.text:
                    pieces.append(mt.text)
        else:
            # fallback: any w:t deeper in this node
            for t in child.findall('.//w:t', namespaces=NSMAP):
                if t.text:
                    pieces.append(t.text)
    return "".join([p if p is not None else "" for p in pieces]).strip()

def normalize_text(s):
    """Unicode NFC and collapse whitespace, remove zero-width chars."""
    if s is None:
        return ""
    s = unicodedata.normalize("NFC", s)
    s = s.replace('\u200b', '').replace('\u200c', '').replace('\u200d', '')
    s = re.sub(r'\r\n|\r', '\n', s)
    s = re.sub(r'\n\s+\n', '\n\n', s)
    s = re.sub(r'[ \t]+', ' ', s)
    return s.strip()

# ----------------------
# Parsing algorithm
# ----------------------

SEPARATOR_RE = re.compile(r'^[\-\—\–\*\_]{1,}\s*$')

HEADING_PATTERNS = [
    r'^unique questions',
    r'^unique questions with answers',
    r'^paper\s*\d+',
    r'^selected questions',
    r'^selected questions with answers',
    r'^questions? on',
]

def is_heading_line(line: str) -> bool:
    if not line:
        return False
    ln = line.strip().lower()
    for pat in HEADING_PATTERNS:
        if re.match(pat, ln):
            return True
    if 'unique questions' in ln or 'answers and explanations' in ln:
        return True
    if re.match(r'^paper\s*\d+\b', ln):
        return True
    return False

def group_paragraphs_into_blocks(paragraphs):
    blocks = []
    curr = []
    for p in paragraphs:
        if p is None:
            continue
        if p.strip() == "" or SEPARATOR_RE.match(p.strip()):
            if curr:
                blocks.append(curr)
                curr = []
            continue
        curr.append(p)
    if curr:
        blocks.append(curr)
    return blocks

# Core block parsing
def parse_block(block, idx):
    """
    Parse a block (list of paragraph strings) into:
    { question, options[4], answer ('a'..'d'), assumed(bool), explanation, raw_block }
    Returns None if block is considered heading/noise.
    """
    paras = [p for p in block if p and p.strip()!='']
    if not paras:
        return None
    # remove leading headings if present
    while paras and is_heading_line(paras[0]):
        paras.pop(0)
    if not paras:
        return None

    raw = " ".join(paras).strip()
    debug = {"block_index": idx, "raw_block_preview": raw[:240]}

    # extract explanation at end, if exists
    explanation = ""
    m_expl = re.search(r'(?i)\b(Explanation|Solution|Explanatory)\b\s*[:\-]?\s*(.*)$', raw, flags=re.S)
    if m_expl:
        explanation = m_expl.group(2).strip()
        raw = raw[:m_expl.start()].strip()

    # extract answer if present anywhere (prefer before options split)
    raw_answer_text = None
    answer_letter = None
    m_ans = re.search(r'(?i)\b(Answer|Ans|Correct|Key|Correct option)\b\s*[:\-]?\s*([^\n\r]*)', raw)
    if m_ans:
        raw_answer_text = m_ans.group(2).strip()
        raw = raw[:m_ans.start()].strip()
        m_letter = re.search(r'([A-Da-d])', raw_answer_text)
        if m_letter:
            answer_letter = m_letter.group(1).lower()

    # split question vs options
    question_text = raw
    options_part = ""
    m_options_token = re.search(r'(?i)\bOptions?\b\s*[:\-]?\s*(.*)$', raw, flags=re.S)
    if m_options_token:
        question_text = raw[:m_options_token.start()].strip()
        options_part = m_options_token.group(1).strip()
    else:
        # attempt split before first labeled option (a) or a. etc
        sp = re.split(r'(?=(?:\(|\[)?[A-Da-d][\)\].]?\s+)', raw, maxsplit=1)
        if len(sp) == 2:
            question_text, options_part = sp[0].strip(), sp[1].strip()
        else:
            # maybe options are on separate paragraphs - gather lines starting with a., (a) etc
            separate_opts = []
            for p in paras[1:]:
                m = re.match(r'^\s*[\(\[]?([A-Da-d])[\)\].]?\s*(.+)', p)
                if m:
                    separate_opts.append((m.group(1).lower(), m.group(2).strip()))
            if separate_opts:
                # build options_part by concatenating labeled lines
                options_part = " ".join([f"({lab}) {txt}" for lab,txt in separate_opts])
                # question_text remains the first paragraph
                question_text = paras[0].strip()

    # extract pairs label -> text using robust regex
    opt_pairs = re.findall(
        r'[\(\[]?([A-Da-d])[\)\].]?\s*'          # label
        r'([^(\(\[]+?)'                          # option text (lazy)
        r'(?=(?:[\(\[]?[A-Da-d][\)\].]?|\Z))',   # until next label or end
        options_part, flags=re.S
    )
    opts = []
    if opt_pairs:
        # sort by label order a..d
        label_map = {lab.lower(): txt.strip() for lab, txt in opt_pairs}
        for label in ['a','b','c','d']:
            opts.append(normalize_text(label_map.get(label, "")))
    else:
        # fallback: try to find inline 'a) text b) text' by splitting tokens
        inline = re.split(r'[\s]*[A-Da-d][\)\.\]]\s*', options_part)
        inline = [normalize_text(x) for x in inline if normalize_text(x)]
        if inline:
            # note: the split yields leading prefix before first label; to be safe, take last 4
            if len(inline) >= 4:
                opts = inline[:4]
            else:
                opts = inline + [""] * (4 - len(inline))
        else:
            opts = ["", "", "", ""]

    # ensure 4 options
    if len(opts) < 4:
        opts += [""] * (4 - len(opts))
    if len(opts) > 4:
        opts = opts[:4]

    # try to resolve answer by matching raw_answer_text against options if letter unknown
    assumed = False
    if not answer_letter and raw_answer_text:
        ra = raw_answer_text.lower()
        for i, o in enumerate(opts):
            if o and o.lower() in ra:
                answer_letter = 'abcd'[i]
                break

    if not answer_letter:
        # fallback choose 'a' and mark assumed
        answer_letter = 'a'
        assumed = True

    parsed = {
        'question': normalize_text(question_text),
        'options': opts,
        'answer': answer_letter,
        'assumed': assumed,
        'explanation': normalize_text(explanation),
        'raw_block': normalize_text(raw),
        'debug': {
            'block_idx': idx,
            'raw_preview': raw[:220],
            'found_options': opts,
            'raw_answer_text': raw_answer_text,
            'final_answer': answer_letter,
            'assumed': assumed
        }
    }

    # If parsed question empty and options empty treat as not a question
    if (not parsed['question']) and all(o == "" for o in parsed['options']):
        return None
    return parsed

def parse_docx_to_questions(input_path, write_debug_log=True):
    """
    Parses the input docx and returns a list of parsed question dicts.
    Also writes debug_log.jsonl if write_debug_log True.
    """
    doc = Document(input_path)
    paragraphs = [paragraph_full_text(p) for p in doc.paragraphs]
    blocks = group_paragraphs_into_blocks(paragraphs)
    questions = []
    debug_entries = []
    for i, block in enumerate(blocks):
        q = parse_block(block, i)
        if q:
            questions.append(q)
            debug_entries.append(q['debug'])
    if write_debug_log:
        try:
            with open("debug_log.jsonl", "w", encoding="utf-8") as f:
                for e in debug_entries:
                    f.write(json.dumps(e, ensure_ascii=False) + "\n")
        except Exception as e:
            # ignore write debug errors; not critical
            pass
    return questions

# ----------------------
# Output builder
# ----------------------

def write_output_docx(questions, output_path):
    """
    Write the exact required table per question:
    8 rows x 3 cols:
    Row0: cell(0,0) = 'Question', cell(0,1)+cell(0,2) merged -> question text
    Row1: cell(1,0) = 'Type', merge cell(1,1..2) -> 'multiple_choice'
    Row2-5: each row: 'Option' | option text | correct/incorrect
    Row6: 'Solution' | merge cell(6,1..2) -> explanation
    Row7: 'Marks' | '1' | '0'
    """
    doc = Document()
    for q in questions:
        table = doc.add_table(rows=8, cols=3)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass

        # Row0
        table.cell(0,0).text = "Question"
        table.cell(0,1).merge(table.cell(0,2)).text = q.get('question', '')

        # Row1
        table.cell(1,0).text = "Type"
        table.cell(1,1).merge(table.cell(1,2)).text = "multiple_choice"

        # Rows 2-5: Options
        opts = q.get('options', ['','','',''])
        letter = q.get('answer','a').lower()
        idx_map = {'a':0,'b':1,'c':2,'d':3}
        correct_idx = idx_map.get(letter, 0)
        for i in range(4):
            r = 2 + i
            table.cell(r,0).text = "Option"
            table.cell(r,1).text = opts[i] or ""
            table.cell(r,2).text = "correct" if i==correct_idx else "incorrect"

        # Row6 solution
        table.cell(6,0).text = "Solution"
        table.cell(6,1).merge(table.cell(6,2)).text = q.get('explanation','')

        # Row7 marks
        table.cell(7,0).text = "Marks"
        table.cell(7,1).text = "1"
        table.cell(7,2).text = "0"

        # spacing paragraph
        doc.add_paragraph()
    doc.save(output_path)

# --- Conversion Logic End ---

# ------------ QSS (light & dark) ------------
# We keep both in case you want to switch default later.
DARK_QSS = r"""
/* Base background and font */
QWidget { background: #0b0f13; color: #e6eef8; font-family: "Segoe UI", Arial, sans-serif; }

/* Topbar */
#topbar { background: #0f1720; min-height: 64px; max-height: 64px; }
#app_name { color: #ffffff; font-weight: 800; font-size: 32px; padding-left:6px; }

/* Card */
#card { background: #0f1728; border-radius: 12px; }

/* Labels & inputs */
QLabel { background: transparent; color: #cbd5e1; font-size: 15px; }
QLabel#success_msg { color: #22c55e; font-size: 16px; }
QLineEdit { background: transparent; border: none; color: #e6eef8; font-size: 15px; padding: 6px 8px; border-bottom: 1px solid #1f2937; }

/* Thin bottom border containers to mimic modern underline input */
#input_line, #output_line {
  padding-bottom: 8px;
}

/* Pill buttons: using dynamic property pclass="pill" */
QPushButton[pclass="pill"] {
  background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #0a84ff, stop:1 #0a84ff);
  color: white;
  border-radius: 999px;
  min-width: 110px;
  min-height: 36px;
  font-size: 16px;
  padding: 6px 16px;
  border: none;
}
QPushButton[pclass="pill"]:hover { background: #0066cc; }

/* Convert button slightly smaller height than before for your request */
QPushButton#convert_btn[pclass="pill"] {
  min-height: 40px;
  font-size: 16px;
  padding: 8px 18px;
}

/* Theme toggle button */
QPushButton#theme_btn {
  background: rgba(255,255,255,0.04);
  color: #d1d5db;
  border: none;
  min-width: 40px;
  min-height: 40px;
  border-radius: 8px;
}

/* Reset button */
QPushButton#reset_btn {
  background: rgba(255,255,255,0.04);
  color: #d1d5db;
  border: none;
  min-width: 80px;
  min-height: 40px;
  border-radius: 8px;
  font-size: 14px;
}

/* Toast */
#toast {
  background: rgba(255,255,255,0.06);
  color: #e6eef8;
  padding: 10px 14px;
  border-radius: 10px;
}

/* Remove frame borders globally (we use QFrame as container) */
QFrame { border: none; }
"""

LIGHT_QSS = r"""
QWidget { background: #f5f5f5; color: #0f1720; font-family: "Segoe UI", Arial, sans-serif; }

/* Topbar */
#topbar { background: #111827; min-height: 64px; max-height: 64px; }
#app_name { color: #ffffff; font-weight: 800; font-size: 32px; padding-left:6px; }

/* Card */
#card { background: #ffffff; border-radius: 12px; }

/* Labels & inputs */
QLabel { background: transparent; color: #374151; font-size: 15px; }
QLabel#success_msg { color: #22c55e; font-size: 16px; }
QLineEdit { background: transparent; border: none; color: #0f1720; font-size: 15px; padding: 6px 8px; border-bottom: 1px solid #e6e6e6; }

/* underline style */
#input_line, #output_line {
  padding-bottom: 8px;
}

/* pill */
QPushButton[pclass="pill"] {
  background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #007bff, stop:1 #007bff);
  color: white;
  border-radius: 999px;
  min-width: 110px;
  min-height: 36px;
  font-size: 16px;
  padding: 6px 16px;
  border: none;
}
QPushButton[pclass="pill"]:hover { background: #0056b3; }

QPushButton#convert_btn[pclass="pill"] {
  min-height: 40px;
  font-size: 16px;
  padding: 8px 18px;
}

/* Theme toggle */
QPushButton#theme_btn {
  background: rgba(0,0,0,0.06);
  color: #111827;
  border: none;
  min-width: 40px;
  min-height: 40px;
  border-radius: 8px;
}

/* Reset button */
QPushButton#reset_btn {
  background: rgba(0,0,0,0.15); /* Slightly darker background for better contrast */
  color: #ffffff; /* White text for high contrast against the background */
  border: 1px solid #d1d5db; /* Added border for definition */
  min-width: 80px;
  min-height: 40px;
  border-radius: 8px;
  font-size: 14px;
  padding: 0 10px; /* Added padding to ensure text fits well */
}
QPushButton#reset_btn:hover {
  background: rgba(0,0,0,0.25); /* Darker on hover for feedback */
  color: #ffffff;
}

/* Toast */
#toast { background: rgba(0,0,0,0.65); color: #fff; padding: 10px 14px; border-radius: 10px; }
QFrame { border: none; }
"""

# -------------------- Main Window --------------------
class QuizFormatterMain(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("QuizFormatter")
        self.setMinimumSize(880, 560)

        # central layout
        central = QWidget()
        self.setCentralWidget(central)
        main_v = QVBoxLayout(central)
        main_v.setContentsMargins(0, 0, 0, 14)
        main_v.setSpacing(0)

        # ----- Topbar (compact) -----
        topbar = QWidget(objectName="topbar")
        topbar_layout = QHBoxLayout(topbar)
        topbar_layout.setContentsMargins(18, 10, 18, 10)
        topbar_layout.setSpacing(8)

        self.app_name = QLabel("QuizFormatter", objectName="app_name")
        # fallback font; QSS will set size, but keep weight
        self.app_name.setFont(QFont("Segoe UI", 28, QFont.Weight.Bold))
        topbar_layout.addWidget(self.app_name, alignment=Qt.AlignLeft | Qt.AlignVCenter)

        topbar_layout.addStretch()

        # Reset button — added before theme button
        self.reset_btn = QPushButton("Reset", objectName="reset_btn")
        self.reset_btn.setCursor(Qt.PointingHandCursor)
        self.reset_btn.setFixedSize(80, 40)
        self.reset_btn.clicked.connect(self.reset_app)
        self.reset_btn.setToolTip("Reset to starting state")
        topbar_layout.addWidget(self.reset_btn, alignment=Qt.AlignRight | Qt.AlignVCenter)

        # Theme toggle button (shows moon or sun) — more intuitive than gear
        self.theme_btn = QPushButton("🌙", objectName="theme_btn")
        self.theme_btn.setCursor(Qt.PointingHandCursor)
        self.theme_btn.setFixedSize(40, 40)
        self.theme_btn.clicked.connect(self.toggle_theme)
        topbar_layout.addWidget(self.theme_btn, alignment=Qt.AlignRight | Qt.AlignVCenter)

        main_v.addWidget(topbar)

        # ----- Content area — centered card -----
        content = QWidget()
        content_layout = QVBoxLayout(content)
        content_layout.setContentsMargins(0, 18, 0, 18)
        content_layout.setSpacing(0)

        content_layout.addStretch()

        # center horizontally
        hwrap = QHBoxLayout()
        hwrap.addStretch()

        card = QFrame(objectName="card")
        card.setFixedWidth(480)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(32, 26, 32, 26)
        card_layout.setSpacing(16)

        # Input
        lbl_input = QLabel("Input")
        lbl_input.setFont(QFont("Segoe UI", 13))
        card_layout.addWidget(lbl_input)

        input_line = QWidget(objectName="input_line")
        input_layout = QHBoxLayout(input_line)
        input_layout.setContentsMargins(6, 0, 6, 6)
        input_layout.setSpacing(12)

        self.input_edit = QLineEdit(placeholderText="Choose file...")
        self.input_edit.setFixedHeight(36)
        self.input_edit.setReadOnly(True)
        self.input_edit.setFont(QFont("Segoe UI", 13))
        self.input_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        self.browse_btn = QPushButton("Browse")
        self.browse_btn.setProperty("pclass", "pill")
        self.browse_btn.setCursor(Qt.PointingHandCursor)
        self.browse_btn.setFixedHeight(36)
        self.browse_btn.setFixedWidth(120)
        self.browse_btn.setFlat(True)
        self.browse_btn.clicked.connect(self.browse_file)

        input_layout.addWidget(self.input_edit)
        input_layout.addWidget(self.browse_btn)
        card_layout.addWidget(input_line)

        # Output
        lbl_output = QLabel("Output")
        lbl_output.setFont(QFont("Segoe UI", 13))
        card_layout.addWidget(lbl_output)

        output_line = QWidget(objectName="output_line")
        output_layout = QHBoxLayout(output_line)
        output_layout.setContentsMargins(6, 0, 6, 6)
        output_layout.setSpacing(12)

        self.output_edit = QLineEdit(placeholderText="Save as...")
        self.output_edit.setFixedHeight(36)
        self.output_edit.setFont(QFont("Segoe UI", 13))
        self.output_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        self.save_btn = QPushButton("Save")
        self.save_btn.setProperty("pclass", "pill")
        self.save_btn.setCursor(Qt.PointingHandCursor)
        self.save_btn.setFixedHeight(36)
        self.save_btn.setFixedWidth(120)
        self.save_btn.setFlat(True)
        self.save_btn.clicked.connect(self.save_as)

        output_layout.addWidget(self.output_edit)
        output_layout.addWidget(self.save_btn)
        card_layout.addWidget(output_line)

        # Convert button — full width, slightly reduced height, pill
        self.convert_btn = QPushButton("Convert", objectName="convert_btn")
        self.convert_btn.setProperty("pclass", "pill")
        self.convert_btn.setCursor(Qt.PointingHandCursor)
        self.convert_btn.setFixedHeight(40)
        self.convert_btn.setFlat(True)
        self.convert_btn.clicked.connect(self.convert)
        card_layout.addSpacing(6)
        card_layout.addWidget(self.convert_btn)

        # Success message label, initially hidden
        self.success_label = QLabel("✅ Successfully Converted", objectName="success_msg")
        self.success_label.setAlignment(Qt.AlignCenter)
        self.success_label.setVisible(False)
        card_layout.addSpacing(6)
        card_layout.addWidget(self.success_label)

        card_layout.addStretch()
        hwrap.addWidget(card)
        hwrap.addStretch()
        content_layout.addLayout(hwrap)

        content_layout.addStretch()
        main_v.addWidget(content)

        # Toast
        self.toast = QLabel("", objectName="toast")
        self.toast.setVisible(False)
        self.toast.setAlignment(Qt.AlignCenter)
        toast_wrap = QHBoxLayout()
        toast_wrap.addStretch()
        toast_wrap.addWidget(self.toast)
        toast_wrap.addStretch()
        main_v.addLayout(toast_wrap)

        # state
        self.current_input_path = None
        self.output_path = None

        # default to dark theme as requested
        self.is_dark = True

        # IMPORTANT: use Fusion style so QSS reliably applies across Windows versions
        app = QApplication.instance()
        app.setStyle("Fusion")
        # apply theme (app-level stylesheet) now
        self.apply_theme()

        # shortcuts
        focus_output = QAction(self)
        focus_output.setShortcut(QKeySequence("Ctrl+K"))
        focus_output.triggered.connect(lambda: self.output_edit.setFocus())
        self.addAction(focus_output)

        # ensure theme button icon matches initial theme
        self.update_theme_icon()

    # ---------- Reset app state ----------
    def reset_app(self):
        # Clear input and output fields
        self.input_edit.clear()
        self.output_edit.clear()
        # Reset paths
        self.current_input_path = None
        self.output_path = None
        # Hide success message
        self.success_label.setVisible(False)
        # Optional: reset to dark theme if desired, but keeping current theme
        # Show toast for feedback
        self.show_toast("App reset to starting state", 1400)

    # ---------- Styling helpers ----------
    def apply_theme(self):
        app = QApplication.instance()
        if app is None:
            return
        if self.is_dark:
            app.setStyleSheet(DARK_QSS)
        else:
            app.setStyleSheet(LIGHT_QSS)
        # ensure app-name weight & font stays strong
        self.app_name.setFont(QFont("Segoe UI", 30, QFont.Weight.DemiBold))
        self.update_theme_icon()

    def toggle_theme(self):
        self.is_dark = not self.is_dark
        self.apply_theme()

    def update_theme_icon(self):
        # show moon for dark, sun for light
        if self.is_dark:
            self.theme_btn.setText("☀")  # shows sun icon to indicate "switch to light"
            self.theme_btn.setToolTip("Switch to light theme")
        else:
            self.theme_btn.setText("🌙")  # shows moon icon to indicate "switch to dark"
            self.theme_btn.setToolTip("Switch to dark theme")

    def show_toast(self, text: str, ms: int = 2000):
        self.toast.setText(text)
        self.toast.setVisible(True)
        QTimer.singleShot(ms, lambda: self.toast.setVisible(False))

    # ---------- File actions ----------
    def browse_file(self):
        fname, _ = QFileDialog.getOpenFileName(self, "Select input file", str(Path.home()),
                                              "Word Documents (*.docx);;All Files (*)")
        if not fname:
            return
        self.current_input_path = Path(fname)
        self.input_edit.setText(self.current_input_path.name)
        # suggest output (no forced .json placeholder)
        if not self.output_edit.text().strip():
            self.output_edit.setText(self.current_input_path.stem + "_Formatted")
        self.show_toast(f"Loaded: {self.current_input_path.name}", 1400)

    def save_as(self):
        default_name = (self.output_edit.text().strip() or
                        (self.current_input_path.stem + "_Formatted" if self.current_input_path else "formatted-quiz"))
        fname, _ = QFileDialog.getSaveFileName(self, "Save output as", str(Path.home() / default_name),
                                              "Word Documents (*.docx);;All Files (*)")
        if not fname:
            return
        self.output_path = Path(fname)
        # show filename only, like the web mock
        self.output_edit.setText(self.output_path.name)
        self.show_toast("Save location set", 1200)

    # ---------- Convert (plug backend here) ----------
    def convert(self):
        # Hide success message at the start of conversion
        self.success_label.setVisible(False)

        if not self.current_input_path or not self.current_input_path.exists():
            QMessageBox.warning(self, "No input", "Please choose a valid input file first.")
            return

        out_name_text = self.output_edit.text().strip()
        if self.output_path:
            out_path = self.output_path
        else:
            out_path = self.current_input_path.parent / (out_name_text or (self.current_input_path.stem + "_Formatted.docx"))

        try:
            questions = parse_docx_to_questions(self.current_input_path, write_debug_log=True)
        except Exception as ex:
            QMessageBox.critical(self, "Parse error", f"Failed to parse input file:\n{ex}")
            return

        # Check for warnings
        warnings = []
        for i, q in enumerate(questions):
            if q.get('assumed'):
                warnings.append(f"Q{i+1}: assumed answer (A).")
            if any(o.strip() == "" for o in q['options']):
                warnings.append(f"Q{i+1}: fewer than 4 options (padded empty).")
        if warnings:
            msg = "Parser made the following assumptions or found issues:\n\n" + "\n".join(warnings[:10]) + ("\n\n(Showing first 10)\n\nProceed with export?" if len(warnings) > 10 else "\n\nProceed with export?")
            reply = QMessageBox.question(self, "Parsing Warnings", msg, QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.No:
                self.show_toast("Export canceled due to warnings.", 1800)
                return

        try:
            write_output_docx(questions, out_path)
            self.show_toast(f"Formatted and saved: {out_path.name}", 1800)
            # Show success message after successful conversion
            self.success_label.setVisible(True)
        except Exception as ex:
            QMessageBox.critical(self, "Save error", f"Failed to save output file:\n{ex}")

# ------------ entrypoint ------------
def main():
    app = QApplication(sys.argv)
    app.setApplicationName("QuizFormatter")
    # ensure consistent QSS behavior on Windows
    app.setStyle("Fusion")

    window = QuizFormatterMain()
    window.show()

    # center the window
    screen = app.primaryScreen().availableGeometry()
    geom = window.geometry()
    window.move((screen.width() - geom.width()) // 2, (screen.height() - geom.height()) // 2)

    sys.exit(app.exec())


if __name__ == "__main__":
    main()