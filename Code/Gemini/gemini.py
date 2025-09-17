# quiz_formatter_app_windows_ready_with_gemini.py
# PySide6 desktop UI with Google Gemini API integration for robust parsing.
# pip install PySide6 python-docx lxml google-generativeai python-dotenv

import sys, os, json, re
from pathlib import Path
from docx import Document
from PySide6.QtCore import Qt, QTimer, QThread, Signal
from PySide6.QtGui import QFont, QKeySequence, QAction
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QLineEdit, QPushButton,
    QHBoxLayout, QVBoxLayout, QFileDialog, QMessageBox, QFrame, QSizePolicy,
    QProgressDialog
)
from dotenv import load_dotenv
import google.generativeai as genai

# Load environment variables from .env file
load_dotenv()

# --- Gemini API Logic Start ---
class GeminiParser(QThread):
    # Signals for communication with the main thread
    parsing_finished = Signal(list, list) # list of questions, list of warnings
    parsing_error = Signal(str)
    progress_updated = Signal(int, str)

    def __init__(self, full_text):
        super().__init__()
        self.full_text = full_text
        self.questions = []
        self.warnings = []
        self._is_running = True

    def run(self):
        try:
            api_key = os.getenv("GOOGLE_API_KEY")
            if not api_key:
                raise ValueError("GOOGLE_API_KEY not found. Please set it in a .env file.")
            
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')

            # Split the document content by separator for individual processing
            blocks = re.split(r'—+', self.full_text)
            
            # The prompt is key to robust parsing
            prompt_template = """
            Analyze the following question block and extract the question, four options, the correct answer, and the explanation.
            The correct answer should be a letter (a, b, c, or d). If the options are not labeled, assume the order is a, b, c, d.
            If the answer is not explicitly stated, assume it's the first option.
            
            Format the output as a JSON array of a single object, like this:
            [
              {{
                "question": "...",
                "options": ["...", "...", "...", "..."],
                "answer": "a",
                "explanation": "..."
              }}
            ]
            
            Ensure all text, including special characters and formulas, is preserved exactly as it appears in the input.
            
            Question Block:
            {block_text}
            """
            
            # Process each block one by one to avoid exceeding token limits for large docs
            for i, block_text in enumerate(blocks):
                if not self._is_running:
                    return
                
                block_text = block_text.strip()
                if not block_text:
                    continue

                self.progress_updated.emit(i + 1, f"Parsing question {i + 1}...")

                prompt = prompt_template.format(block_text=block_text)
                
                # Use a try-except block for each API call to handle potential errors
                try:
                    response = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(response_mime_type="application/json"))
                    data = json.loads(response.text)
                    
                    if data and isinstance(data, list) and len(data) > 0:
                        parsed_data = data[0]
                        
                        # Add a warning if the answer was assumed
                        if parsed_data.get('assumed_answer', False):
                            self.warnings.append(f"Q{i+1}: Answer was not found and was assumed to be 'a'.")

                        # Ensure a complete set of options
                        if len(parsed_data.get('options', [])) < 4:
                            self.warnings.append(f"Q{i+1}: Fewer than 4 options were found.")
                            while len(parsed_data['options']) < 4:
                                parsed_data['options'].append("")
                        
                        self.questions.append(parsed_data)
                    else:
                        self.warnings.append(f"Q{i+1}: Failed to parse block. AI returned empty or invalid JSON.")

                except Exception as e:
                    self.warnings.append(f"Q{i+1}: AI parsing error - {e}")

            self.parsing_finished.emit(self.questions, self.warnings)
            
        except Exception as e:
            self.parsing_error.emit(str(e))

    def stop(self):
        self._is_running = False

# --- Gemini API Logic End ---

# --- Document I/O Start ---
def paragraph_full_text(paragraph):
    # This helper function is still needed to extract raw text
    return "".join(t.text for t in paragraph.runs).strip()

def write_output_docx(questions, output_path):
    doc = Document()
    for q in questions:
        table = doc.add_table(rows=8, cols=3)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass

        table.cell(0, 0).text = "Question"
        table.cell(0, 1).merge(table.cell(0, 2)).text = q.get('question', '')

        table.cell(1, 0).text = "Type"
        table.cell(1, 1).merge(table.cell(1, 2)).text = "multiple_choice"

        opts = q.get('options', ['', '', '', ''])
        # AI returns a letter, so find its index
        answer_letter = q.get('answer', 'a').lower()
        idx_map = {'a': 0, 'b': 1, 'c': 2, 'd': 3}
        correct_idx = idx_map.get(answer_letter, 0)
        
        for i in range(4):
            r = 2 + i
            table.cell(r, 0).text = "Option"
            table.cell(r, 1).text = opts[i] or ""
            table.cell(r, 2).text = "correct" if i == correct_idx else "incorrect"

        table.cell(6, 0).text = "Solution"
        table.cell(6, 1).merge(table.cell(6, 2)).text = q.get('explanation', '')

        table.cell(7, 0).text = "Marks"
        table.cell(7, 1).text = "1"
        table.cell(7, 2).text = "0"
        
        doc.add_paragraph()
    doc.save(output_path)
# --- Document I/O End ---

# --- UI and Main Application ---
class QuizFormatterMain(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("QuizFormatter")
        self.setMinimumSize(880, 560)
        
        self.current_input_path = None
        self.output_path = None
        self.is_dark = True
        self.parser_thread = None
        
        # --- UI Setup ---
        central = QWidget()
        self.setCentralWidget(central)
        main_v = QVBoxLayout(central)
        main_v.setContentsMargins(0, 0, 0, 14)
        main_v.setSpacing(0)
        topbar = QWidget(objectName="topbar")
        topbar_layout = QHBoxLayout(topbar)
        topbar_layout.setContentsMargins(18, 10, 18, 10)
        topbar_layout.setSpacing(8)
        self.app_name = QLabel("QuizFormatter", objectName="app_name")
        self.app_name.setFont(QFont("Segoe UI", 28, QFont.Weight.Bold))
        topbar_layout.addWidget(self.app_name, alignment=Qt.AlignLeft | Qt.AlignVCenter)
        topbar_layout.addStretch()
        self.reset_btn = QPushButton("Reset", objectName="reset_btn")
        self.reset_btn.setCursor(Qt.PointingHandCursor)
        self.reset_btn.setFixedSize(80, 40)
        self.reset_btn.clicked.connect(self.reset_app)
        self.reset_btn.setToolTip("Reset to starting state")
        topbar_layout.addWidget(self.reset_btn, alignment=Qt.AlignRight | Qt.AlignVCenter)
        self.theme_btn = QPushButton("🌙", objectName="theme_btn")
        self.theme_btn.setCursor(Qt.PointingHandCursor)
        self.theme_btn.setFixedSize(40, 40)
        self.theme_btn.clicked.connect(self.toggle_theme)
        topbar_layout.addWidget(self.theme_btn, alignment=Qt.AlignRight | Qt.AlignVCenter)
        main_v.addWidget(topbar)
        content = QWidget()
        content_layout = QVBoxLayout(content)
        content_layout.setContentsMargins(0, 18, 0, 18)
        content_layout.setSpacing(0)
        content_layout.addStretch()
        hwrap = QHBoxLayout()
        hwrap.addStretch()
        card = QFrame(objectName="card")
        card.setFixedWidth(480)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(32, 26, 32, 26)
        card_layout.setSpacing(16)
        lbl_input = QLabel("Input")
        lbl_input.setFont(QFont("Segoe UI", 13))
        card_layout.addWidget(lbl_input)
        input_line = QWidget(objectName="input_line")
        input_layout = QHBoxLayout(input_line)
        input_layout.setContentsMargins(6, 0, 6, 6)
        input_layout.setSpacing(12)
        self.input_edit = QLineEdit(placeholderText="Choose file...")
        self.input_edit.setFixedHeight(36)
        self.input_edit.setReadOnly(True)
        self.input_edit.setFont(QFont("Segoe UI", 13))
        self.input_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.browse_btn = QPushButton("Browse")
        self.browse_btn.setProperty("pclass", "pill")
        self.browse_btn.setCursor(Qt.PointingHandCursor)
        self.browse_btn.setFixedHeight(36)
        self.browse_btn.setFixedWidth(120)
        self.browse_btn.setFlat(True)
        self.browse_btn.clicked.connect(self.browse_file)
        input_layout.addWidget(self.input_edit)
        input_layout.addWidget(self.browse_btn)
        card_layout.addWidget(input_line)
        lbl_output = QLabel("Output")
        lbl_output.setFont(QFont("Segoe UI", 13))
        card_layout.addWidget(lbl_output)
        output_line = QWidget(objectName="output_line")
        output_layout = QHBoxLayout(output_line)
        output_layout.setContentsMargins(6, 0, 6, 6)
        output_layout.setSpacing(12)
        self.output_edit = QLineEdit(placeholderText="Save as...")
        self.output_edit.setFixedHeight(36)
        self.output_edit.setFont(QFont("Segoe UI", 13))
        self.output_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.save_btn = QPushButton("Save")
        self.save_btn.setProperty("pclass", "pill")
        self.save_btn.setCursor(Qt.PointingHandCursor)
        self.save_btn.setFixedHeight(36)
        self.save_btn.setFixedWidth(120)
        self.save_btn.setFlat(True)
        self.save_btn.clicked.connect(self.save_as)
        output_layout.addWidget(self.output_edit)
        output_layout.addWidget(self.save_btn)
        card_layout.addWidget(output_line)
        self.convert_btn = QPushButton("Convert", objectName="convert_btn")
        self.convert_btn.setProperty("pclass", "pill")
        self.convert_btn.setCursor(Qt.PointingHandCursor)
        self.convert_btn.setFixedHeight(40)
        self.convert_btn.setFlat(True)
        self.convert_btn.clicked.connect(self.convert)
        card_layout.addSpacing(6)
        card_layout.addWidget(self.convert_btn)
        self.success_label = QLabel("✅ Successfully Converted", objectName="success_msg")
        self.success_label.setAlignment(Qt.AlignCenter)
        self.success_label.setVisible(False)
        card_layout.addSpacing(6)
        card_layout.addWidget(self.success_label)
        card_layout.addStretch()
        hwrap.addWidget(card)
        hwrap.addStretch()
        content_layout.addLayout(hwrap)
        content_layout.addStretch()
        main_v.addWidget(content)
        self.toast = QLabel("", objectName="toast")
        self.toast.setVisible(False)
        self.toast.setAlignment(Qt.AlignCenter)
        toast_wrap = QHBoxLayout()
        toast_wrap.addStretch()
        toast_wrap.addWidget(self.toast)
        toast_wrap.addStretch()
        main_v.addLayout(toast_wrap)

        app = QApplication.instance()
        app.setStyle("Fusion")
        self.apply_theme()
        focus_output = QAction(self)
        focus_output.setShortcut(QKeySequence("Ctrl+K"))
        focus_output.triggered.connect(lambda: self.output_edit.setFocus())
        self.addAction(focus_output)
        self.update_theme_icon()

    def reset_app(self):
        self.input_edit.clear()
        self.output_edit.clear()
        self.current_input_path = None
        self.output_path = None
        self.success_label.setVisible(False)
        self.show_toast("App reset to starting state", 1400)

    def apply_theme(self):
        app = QApplication.instance()
        if app is None: return
        app.setStyleSheet(DARK_QSS if self.is_dark else LIGHT_QSS)
        self.app_name.setFont(QFont("Segoe UI", 30, QFont.Weight.DemiBold))
        self.update_theme_icon()

    def toggle_theme(self):
        self.is_dark = not self.is_dark
        self.apply_theme()

    def update_theme_icon(self):
        if self.is_dark:
            self.theme_btn.setText("☀")
            self.theme_btn.setToolTip("Switch to light theme")
        else:
            self.theme_btn.setText("🌙")
            self.theme_btn.setToolTip("Switch to dark theme")

    def show_toast(self, text: str, ms: int = 2000):
        self.toast.setText(text)
        self.toast.setVisible(True)
        QTimer.singleShot(ms, lambda: self.toast.setVisible(False))

    def browse_file(self):
        fname, _ = QFileDialog.getOpenFileName(self, "Select input file", str(Path.home()),
                                               "Word Documents (*.docx);;All Files (*)")
        if not fname: return
        self.current_input_path = Path(fname)
        self.input_edit.setText(self.current_input_path.name)
        if not self.output_edit.text().strip():
            self.output_edit.setText(self.current_input_path.stem + "_Formatted")
        self.show_toast(f"Loaded: {self.current_input_path.name}", 1400)

    def save_as(self):
        default_name = (self.output_edit.text().strip() or
                        (self.current_input_path.stem + "_Formatted" if self.current_input_path else "formatted-quiz"))
        fname, _ = QFileDialog.getSaveFileName(self, "Save output as", str(Path.home() / default_name),
                                               "Word Documents (*.docx);;All Files (*)")
        if not fname: return
        self.output_path = Path(fname)
        self.output_edit.setText(self.output_path.name)
        self.show_toast("Save location set", 1200)

    def convert(self):
        self.success_label.setVisible(False)
        if not self.current_input_path or not self.current_input_path.exists():
            QMessageBox.warning(self, "No input", "Please choose a valid input file first.")
            return

        out_name_text = self.output_edit.text().strip()
        if self.output_path:
            out_path = self.output_path
        else:
            out_path = self.current_input_path.parent / (out_name_text or (self.current_input_path.stem + "_Formatted.docx"))

        try:
            doc = Document(self.current_input_path)
            full_text = "\n".join([paragraph_full_text(p) for p in doc.paragraphs])
            
            self.progress_dialog = QProgressDialog("Parsing with AI...", "Cancel", 0, len(re.split(r'—+', full_text)), self)
            self.progress_dialog.setWindowTitle("Converting...")
            self.progress_dialog.setWindowModality(Qt.WindowModal)
            self.progress_dialog.show()
            
            self.parser_thread = GeminiParser(full_text)
            self.parser_thread.parsing_finished.connect(self.on_parsing_finished)
            self.parser_thread.parsing_error.connect(self.on_parsing_error)
            self.parser_thread.progress_updated.connect(self.on_progress_updated)
            self.progress_dialog.canceled.connect(self.parser_thread.stop)
            
            self.parser_thread.start()

        except Exception as ex:
            QMessageBox.critical(self, "Error", f"Failed to start conversion:\n{ex}")
    
    def on_progress_updated(self, value, message):
        self.progress_dialog.setValue(value)
        self.progress_dialog.setLabelText(message)

    def on_parsing_finished(self, questions, warnings):
        self.progress_dialog.close()
        self.parser_thread = None
        
        if warnings:
            msg = "AI Parser made the following assumptions or found issues:\n\n" + "\n".join(warnings[:10])
            if len(warnings) > 10:
                msg += "\n\n(Showing first 10)\n\nProceed with export?"
            else:
                msg += "\n\nProceed with export?"
            reply = QMessageBox.question(self, "AI Parsing Warnings", msg, QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.No:
                self.show_toast("Export canceled due to warnings.", 1800)
                return

        try:
            write_output_docx(questions, self.output_path)
            self.show_toast(f"Formatted and saved: {self.output_path.name}", 1800)
            self.success_label.setVisible(True)
        except Exception as ex:
            QMessageBox.critical(self, "Save error", f"Failed to save output file:\n{ex}")

    def on_parsing_error(self, error_message):
        self.progress_dialog.close()
        self.parser_thread = None
        QMessageBox.critical(self, "AI Service Error", f"An error occurred with the AI API:\n{error_message}")

# --- QSS and Main Entry Point (unchanged) ---
DARK_QSS = r"""
QWidget { background: #0b0f13; color: #e6eef8; font-family: "Segoe UI", Arial, sans-serif; }
#topbar { background: #0f1720; min-height: 64px; max-height: 64px; }
#app_name { color: #ffffff; font-weight: 800; font-size: 32px; padding-left:6px; }
#card { background: #0f1728; border-radius: 12px; }
QLabel { background: transparent; color: #cbd5e1; font-size: 15px; }
QLabel#success_msg { color: #22c55e; font-size: 16px; }
QLineEdit { background: transparent; border: none; color: #e6eef8; font-size: 15px; padding: 6px 8px; border-bottom: 1px solid #1f2937; }
#input_line, #output_line { padding-bottom: 8px; }
QPushButton[pclass="pill"] {
  background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #0a84ff, stop:1 #0a84ff);
  color: white;
  border-radius: 999px;
  min-width: 110px;
  min-height: 36px;
  font-size: 16px;
  padding: 6px 16px;
  border: none;
}
QPushButton[pclass="pill"]:hover { background: #0066cc; }
QPushButton#convert_btn[pclass="pill"] { min-height: 40px; font-size: 16px; padding: 8px 18px; }
QPushButton#theme_btn {
  background: rgba(255,255,255,0.04);
  color: #d1d5db;
  border: none;
  min-width: 40px;
  min-height: 40px;
  border-radius: 8px;
}
QPushButton#reset_btn {
  background: rgba(255,255,255,0.04);
  color: #d1d5db;
  border: none;
  min-width: 80px;
  min-height: 40px;
  border-radius: 8px;
  font-size: 14px;
}
#toast {
  background: rgba(255,255,255,0.06);
  color: #e6eef8;
  padding: 10px 14px;
  border-radius: 10px;
}
QFrame { border: none; }
"""

LIGHT_QSS = r"""
QWidget { background: #f5f5f5; color: #0f1720; font-family: "Segoe UI", Arial, sans-serif; }
#topbar { background: #111827; min-height: 64px; max-height: 64px; }
#app_name { color: #ffffff; font-weight: 800; font-size: 32px; padding-left:6px; }
#card { background: #ffffff; border-radius: 12px; }
QLabel { background: transparent; color: #374151; font-size: 15px; }
QLabel#success_msg { color: #22c55e; font-size: 16px; }
QLineEdit { background: transparent; border: none; color: #0f1720; font-size: 15px; padding: 6px 8px; border-bottom: 1px solid #e6e6e6; }
#input_line, #output_line { padding-bottom: 8px; }
QPushButton[pclass="pill"] {
  background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #007bff, stop:1 #007bff);
  color: white;
  border-radius: 999px;
  min-width: 110px;
  min-height: 36px;
  font-size: 16px;
  padding: 6px 16px;
  border: none;
}
QPushButton[pclass="pill"]:hover { background: #0056b3; }
QPushButton#convert_btn[pclass="pill"] { min-height: 40px; font-size: 16px; padding: 8px 18px; }
QPushButton#theme_btn {
  background: rgba(0,0,0,0.06);
  color: #111827;
  border: none;
  min-width: 40px;
  min-height: 40px;
  border-radius: 8px;
}
/* Reset button */
QPushButton#reset_btn {
  background: rgba(0,0,0,0.15); /* Slightly darker background for better contrast */
  color: #ffffff; /* White text for high contrast against the background */
  border: 1px solid #d1d5db; /* Added border for definition */
  min-width: 80px;
  min-height: 40px;
  border-radius: 8px;
  font-size: 14px;
  padding: 0 10px; /* Added padding to ensure text fits well */
}
QPushButton#reset_btn:hover {
  background: rgba(0,0,0,0.25); /* Darker on hover for feedback */
  color: #ffffff;
}
#toast { background: rgba(0,0,0,0.65); color: #fff; padding: 10px 14px; border-radius: 10px; }
QFrame { border: none; }
"""

def main():
    app = QApplication(sys.argv)
    app.setApplicationName("QuizFormatter")
    app.setStyle("Fusion")
    window = QuizFormatterMain()
    window.show()
    screen = app.primaryScreen().availableGeometry()
    geom = window.geometry()
    window.move((screen.width() - geom.width()) // 2, (screen.height() - geom.height()) // 2)
    sys.exit(app.exec())

if __name__ == "__main__":
    main()